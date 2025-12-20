# -*- coding: utf-8 -*-
from __future__ import annotations

from io import BytesIO
from typing import Dict, List

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def export_exam_docx(meta: Dict[str, str], questions: List[Dict[str, object]], include_answer_key: bool = True) -> bytes:
    """Xuất docx: KHÔNG tạo 'thang điểm + nhận xét' (theo yêu cầu)."""
    doc = Document()

    title = meta.get("title","ĐỀ KIỂM TRA")
    subtitle = meta.get("subtitle","")
    school = meta.get("school","")
    subject = meta.get("subject","")
    grade = meta.get("grade","")
    term = meta.get("term","")

    p = doc.add_paragraph(school)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(12)

    p = doc.add_paragraph(title)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(16)

    if subtitle.strip():
        p = doc.add_paragraph(subtitle)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = f"Môn: {subject}  |  Lớp: {grade}  |  {term}".strip()
    if info.strip():
        p = doc.add_paragraph(info)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    # Câu hỏi
    for i, q in enumerate(questions, start=1):
        points = q.get("points","")
        level = q.get("level","")
        content = str(q.get("content","")).strip()
        header = f"Câu {i} ({points} đ) - {level}:" if points else f"Câu {i} - {level}:"
        doc.add_paragraph(header).runs[0].bold = True
        doc.add_paragraph(content)

    if include_answer_key:
        doc.add_page_break()
        doc.add_paragraph("ĐÁP ÁN / GỢI Ý").runs[0].bold = True
        for i, q in enumerate(questions, start=1):
            ans = str(q.get("answer","")).strip()
            if not ans:
                # nếu nội dung có dòng "Đáp án:" thì giữ nguyên (GV tự xử lý)
                continue
            doc.add_paragraph(f"Câu {i}: {ans}")

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()
